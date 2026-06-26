import logging
import os
import re
import shutil
from datetime import datetime
from pathlib import Path
from zoneinfo import ZoneInfo

from docx import Document
from openai import APIError, OpenAI, RateLimitError
from pypdf import PdfReader

from .match_score import _tokens

logger = logging.getLogger(__name__)


def resolve_base_resume(path: Path) -> Path:
    """Prefer .docx template (keeps layout); fall back to .pdf."""
    if path.suffix.lower() == ".docx" and path.exists():
        return path
    docx = path.with_suffix(".docx")
    if docx.exists():
        return docx
    if path.exists():
        return path
    raise FileNotFoundError(f"Base resume not found: {path} or {docx}")


def read_resume_text(resume_path: Path) -> str:
    if resume_path.suffix.lower() == ".docx":
        doc = Document(str(resume_path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    reader = PdfReader(str(resume_path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages).strip()


def _slugify(value: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "_", value.lower()).strip("_")
    return slug[:60] or "role"


def _output_folder(output_dir: Path, company: str, job_title: str) -> Path:
    date_prefix = datetime.now(ZoneInfo("America/Chicago")).strftime("%Y-%m-%d")
    folder = output_dir / f"{date_prefix}_{_slugify(company)}_{_slugify(job_title)}"
    folder.mkdir(parents=True, exist_ok=True)
    return folder


def _jd_keywords(job_title: str, job_description: str, limit: int = 12) -> list[str]:
    combined = f"{job_title}\n{job_description}"
    tokens = sorted(_tokens(combined), key=len, reverse=True)
    priority = []
    for token in tokens:
        if token in priority:
            continue
        if len(token) < 3:
            continue
        priority.append(token)
        if len(priority) >= limit:
            break
    return priority


def _paragraph_after_heading(doc: Document, heading: str):
    for index, para in enumerate(doc.paragraphs):
        if para.text.strip().upper() == heading.upper():
            for next_para in doc.paragraphs[index + 1 :]:
                if next_para.text.strip():
                    return next_para
            return None
    return None


def tailor_resume_preserve(
    base_resume_path: Path,
    job_title: str,
    company: str,
    job_description: str,
    output_dir: Path,
) -> tuple[Path, str]:
    """
    Copy your formatted .docx template and only adjust summary + skills text.
    Layout, sections, spacing, and experience bullets stay unchanged.
    """
    base_docx = resolve_base_resume(base_resume_path)
    folder = _output_folder(output_dir, company, job_title)
    dest = folder / "resume.docx"
    shutil.copy2(base_docx, dest)

    doc = Document(dest)
    keywords = _jd_keywords(job_title, job_description)

    summary_para = _paragraph_after_heading(doc, "PROFESSIONAL SUMMARY")
    if summary_para:
        missing = [k for k in keywords[:6] if k.lower() not in summary_para.text.lower()]
        if missing:
            summary_para.text = (
                summary_para.text.rstrip()
                + f" Aligning with {job_title} at {company}: "
                + ", ".join(missing)
                + "."
            )

    skills_para = _paragraph_after_heading(doc, "SKILLS")
    if skills_para:
        existing = skills_para.text
        to_add = [k for k in keywords if k.lower() not in existing.lower()]
        if to_add:
            skills_para.text = existing.rstrip("., ") + ", " + ", ".join(to_add[:8])

    doc.save(dest)
    return dest, "preserved template format; keywords added to summary/skills only"


def tailor_resume(
    base_resume_path: Path,
    job_title: str,
    company: str,
    job_description: str,
    output_dir: Path,
    *,
    mode: str = "preserve",
) -> tuple[Path, str]:
    if mode in ("preserve", "local", "auto"):
        return tailor_resume_preserve(
            base_resume_path, job_title, company, job_description, output_dir
        )

    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if mode == "openai" and api_key:
        logger.warning(
            "openai resume_mode rewrites layout; use resume_mode: preserve to keep your format."
        )
    return tailor_resume_preserve(
        base_resume_path, job_title, company, job_description, output_dir
    )
